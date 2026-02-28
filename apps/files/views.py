from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import UserFile
import os
from django.conf import settings
from django.contrib import messages
from docx import Document
from fpdf import FPDF
from PIL import Image
import pandas as pd
import zipfile
import uuid
from django.contrib import messages
from django.http import Http404, FileResponse
import PyPDF2
from pdf2docx import Converter
import pdfplumber
import openpyxl

@login_required
def upload_file(request):
    if request.method == "POST":
        uploaded_file = request.FILES.get("file")

        if uploaded_file:
            UserFile.objects.create(
                user=request.user,
                original_name=uploaded_file.name,
                file=uploaded_file,
            )

    return redirect("dashboard")


@login_required
def dashboard(request):
    files = UserFile.objects.filter(user=request.user).order_by("-created_at")
    return render(request, "accounts/dashboard.html", {"files": files})


@login_required
def process_tool(request, file_id):
    user_file = get_object_or_404(UserFile, id=file_id, user=request.user)

    if request.method == "POST":
        tool_type = request.POST.get("tool_type")
        
        # Validate tool_type
        valid_tools = {
            ".doc": ["word_to_pdf", "word_to_txt", "compress_word"],
            ".docx": ["word_to_pdf", "word_to_txt", "compress_word"],
            ".pdf": ["pdf_to_word", "pdf_to_txt", "compress_pdf"],
            ".jpg": ["to_jpg", "to_png", "to_gif", "compress_image"],
            ".jpeg": ["to_jpg", "to_png", "to_gif", "compress_image"],
            ".png": ["to_jpg", "to_png", "to_gif", "compress_image"],
            ".gif": ["to_jpg", "to_png", "to_gif", "compress_image"],
            ".zip": ["extract"],
            ".rar": ["extract"],
            ".7z": ["extract"],
            ".csv": ["to_excel"],
            ".xls": ["to_csv"],
            ".xlsx": ["to_csv"],
        }
        
        input_path = user_file.file.path
        name, ext = os.path.splitext(user_file.file.name)
        ext = ext.lower()
        
        # Check if tool_type is valid for this file format
        if ext not in valid_tools or tool_type not in valid_tools[ext]:
            messages.error(request, "Invalid operation for this file type")
            return redirect("dashboard")
        
        user_file.tool_type = tool_type
        
        output_dir = os.path.join(settings.MEDIA_ROOT, 'outputs', f'user_{request.user.id}')
        os.makedirs(output_dir, exist_ok=True)

        output_name = str(uuid.uuid4())
        converted_path = ""

        try:
            # ---------------- WORD FILES (.doc, .docx) ----------------
            if ext in [".doc", ".docx"]:
                doc = Document(input_path)
                
                if tool_type == "word_to_pdf":
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    for para in doc.paragraphs:
                        text = para.text[:1000] if len(para.text) > 1000 else para.text
                        # Handle special characters
                        try:
                            pdf.multi_cell(0, 8, text)
                        except:
                            pdf.multi_cell(0, 8, text.encode('latin-1', 'ignore').decode('latin-1'))
                    converted_path = os.path.join(output_dir, f"{output_name}.pdf")
                    pdf.output(converted_path)
                    messages.success(request, "Word file successfully converted to PDF")

                elif tool_type == "word_to_txt":
                    converted_path = os.path.join(output_dir, f"{output_name}.txt")
                    with open(converted_path, "w", encoding="utf-8") as f:
                        for para in doc.paragraphs:
                            f.write(para.text + "\n")
                    messages.success(request, "Word file successfully converted to TXT")

                elif tool_type == "compress_word":
                    converted_path = os.path.join(output_dir, f"{output_name}.zip")
                    with zipfile.ZipFile(converted_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
                        zipf.write(input_path, arcname=os.path.basename(input_path))
                    messages.success(request, "Word file successfully compressed")

            # ---------------- PDF FILES (.pdf) ----------------
            elif ext == ".pdf":
                
                if tool_type == "pdf_to_word":
                    try:
                        converted_path = os.path.join(output_dir, f"{output_name}.docx")
                        # Method 1: Using pdf2docx (best for formatting)
                        cv = Converter(input_path)
                        cv.convert(converted_path, start=0, end=None)
                        cv.close()
                        messages.success(request, "PDF file successfully converted to Word")
                    except Exception as e:
                        # Method 2: Using PyPDF2 as fallback
                        try:
                            pdf_reader = PyPDF2.PdfReader(input_path)
                            doc = Document()
                            for page_num in range(len(pdf_reader.pages)):
                                page = pdf_reader.pages[page_num]
                                text = page.extract_text()
                                if text:
                                    doc.add_paragraph(text)
                            doc.save(converted_path)
                            messages.success(request, "PDF file successfully converted to Word")
                        except Exception as e2:
                            messages.error(request, f"Error converting PDF to Word: {str(e2)}")
                            return redirect("dashboard")
                    
                elif tool_type == "pdf_to_txt":
                    try:
                        converted_path = os.path.join(output_dir, f"{output_name}.txt")
                        # Use pdfplumber for better text extraction
                        with pdfplumber.open(input_path) as pdf:
                            text = ""
                            for page in pdf.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n\n"
                        
                        with open(converted_path, "w", encoding="utf-8") as f:
                            f.write(text)
                        messages.success(request, "PDF file successfully converted to TXT")
                    except Exception as e:
                        # Fallback to PyPDF2
                        try:
                            pdf_reader = PyPDF2.PdfReader(input_path)
                            text = ""
                            for page_num in range(len(pdf_reader.pages)):
                                page = pdf_reader.pages[page_num]
                                text += page.extract_text() + "\n\n"
                            
                            with open(converted_path, "w", encoding="utf-8") as f:
                                f.write(text)
                            messages.success(request, "PDF file successfully converted to TXT")
                        except Exception as e2:
                            messages.error(request, f"Error converting PDF to TXT: {str(e2)}")
                            return redirect("dashboard")
                    
                elif tool_type == "compress_pdf":
                    converted_path = os.path.join(output_dir, f"{output_name}.zip")
                    with zipfile.ZipFile(converted_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
                        zipf.write(input_path, arcname=os.path.basename(input_path))
                    messages.success(request, "PDF file successfully compressed")

            # ---------------- IMAGE FILES (.jpg, .jpeg, .png, .gif) ----------------
            elif ext in [".jpg", ".jpeg", ".png", ".gif"]:
                try:
                    img = Image.open(input_path)
                    
                    if tool_type == "to_jpg":
                        converted_path = os.path.join(output_dir, f"{output_name}.jpg")
                        if img.mode in ('RGBA', 'LA', 'P'):
                            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                            rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                            rgb_img.save(converted_path, "JPEG", quality=95)
                        else:
                            img.convert("RGB").save(converted_path, "JPEG", quality=95)
                        messages.success(request, "Image successfully converted to JPG")
                        
                    elif tool_type == "to_png":
                        converted_path = os.path.join(output_dir, f"{output_name}.png")
                        img.save(converted_path, "PNG")
                        messages.success(request, "Image successfully converted to PNG")
                        
                    elif tool_type == "to_gif":
                        converted_path = os.path.join(output_dir, f"{output_name}.gif")
                        if img.mode == 'P':
                            img.save(converted_path, "GIF")
                        else:
                            img.convert('P').save(converted_path, "GIF")
                        messages.success(request, "Image successfully converted to GIF")
                        
                    elif tool_type == "compress_image":
                        converted_path = os.path.join(output_dir, f"{output_name}.zip")
                        with zipfile.ZipFile(converted_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
                            zipf.write(input_path, arcname=os.path.basename(input_path))
                        messages.success(request, "Image successfully compressed")
                        
                except Exception as e:
                    messages.error(request, f"Error processing image: {str(e)}")
                    return redirect("dashboard")

            # ---------------- ARCHIVE FILES (.zip, .rar, .7z) ----------------
            elif ext in [".zip", ".rar", ".7z"]:
                if tool_type == "extract":
                    extract_dir = os.path.join(output_dir, f"{output_name}_extracted")
                    os.makedirs(extract_dir, exist_ok=True)
                    
                    # Only ZIP is fully supported
                    if ext == ".zip":
                        try:
                            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                                zip_ref.extractall(extract_dir)
                            
                            # Create zip for download
                            converted_path = os.path.join(output_dir, f"{output_name}.zip")
                            with zipfile.ZipFile(converted_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
                                for root, dirs, files in os.walk(extract_dir):
                                    for file in files:
                                        file_path = os.path.join(root, file)
                                        arcname = os.path.relpath(file_path, extract_dir)
                                        zipf.write(file_path, arcname)
                            
                            messages.success(request, "Archive successfully extracted")
                        except Exception as e:
                            messages.error(request, f"Error extracting archive: {str(e)}")
                            return redirect("dashboard")
                    else:
                        messages.error(request, "Only ZIP format is supported for extraction")
                        return redirect("dashboard")

            # ---------------- CSV FILES (.csv) ----------------
            elif ext == ".csv":
                if tool_type == "to_excel":
                    try:
                        # Try different encodings
                        encodings = ['utf-8', 'latin-1', 'cp1252']
                        df = None
                        
                        for encoding in encodings:
                            try:
                                df = pd.read_csv(input_path, encoding=encoding)
                                break
                            except:
                                continue
                        
                        if df is None:
                            df = pd.read_csv(input_path, encoding='utf-8', errors='ignore')
                        
                        converted_path = os.path.join(output_dir, f"{output_name}.xlsx")
                        df.to_excel(converted_path, index=False, engine='openpyxl')
                        messages.success(request, "CSV file successfully converted to Excel")
                    except Exception as e:
                        messages.error(request, f"Error converting CSV to Excel: {str(e)}")
                        return redirect("dashboard")

            # ---------------- EXCEL FILES (.xls, .xlsx) ----------------
            elif ext in [".xls", ".xlsx"]:
                if tool_type == "to_csv":
                    try:
                        df = pd.read_excel(input_path, engine='openpyxl' if ext == '.xlsx' else None)
                        converted_path = os.path.join(output_dir, f"{output_name}.csv")
                        df.to_csv(converted_path, index=False, encoding="utf-8-sig")
                        messages.success(request, "Excel file successfully converted to CSV")
                    except Exception as e:
                        messages.error(request, f"Error converting Excel to CSV: {str(e)}")
                        return redirect("dashboard")

            # ---------------- Save converted file ----------------
            if converted_path and os.path.isfile(converted_path):
                # Check file size (max 100MB)
                if os.path.getsize(converted_path) > 100 * 1024 * 1024:
                    os.remove(converted_path)
                    messages.error(request, "Output file is too large (max 100MB)")
                    return redirect("dashboard")
                
                relative_path = os.path.relpath(converted_path, settings.MEDIA_ROOT)
                user_file.converted_file.name = relative_path
                user_file.save()
                messages.success(request, "File processed successfully!")
            else:
                messages.error(request, "Error creating output file")

        except Exception as e:
            messages.error(request, f"Processing error: {str(e)}")
            return redirect("dashboard")

    return redirect("dashboard")


@login_required
def download_file(request, file_id):
    user_file = get_object_or_404(UserFile, id=file_id, user=request.user)

    if not user_file.converted_file:
        messages.error(request, "Converted file not found")
        return redirect("dashboard")

    file_path = user_file.converted_file.path
    
    # Additional security check
    if not os.path.exists(file_path):
        messages.error(request, "File not found on server")
        return redirect("dashboard")
    
    # Check if file is in allowed directory
    media_root = os.path.normpath(settings.MEDIA_ROOT)
    file_path_norm = os.path.normpath(file_path)
    if not file_path_norm.startswith(media_root):
        messages.error(request, "Unauthorized access")
        return redirect("dashboard")
    
    try:
        response = FileResponse(open(file_path, 'rb'), as_attachment=True)
        
        # Clean filename
        filename = os.path.basename(file_path)
        # Remove any non-ascii characters for safety
        filename = ''.join(char for char in filename if ord(char) < 128)
        
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Set appropriate Content-Type
        content_types = {
            '.pdf': 'application/pdf',
            '.txt': 'text/plain',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.zip': 'application/zip',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.csv': 'text/csv',
        }
        
        ext = os.path.splitext(filename)[1].lower()
        if ext in content_types:
            response['Content-Type'] = content_types[ext]
        
        return response
        
    except Exception as e:
        messages.error(request, f"Download error: {str(e)}")
        return redirect("dashboard")